import sys
import spacy
import pandas as pd
import matplotlib.pyplot as plt
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QFileDialog, QMessageBox, QTableWidget, QTableWidgetItem,
    QVBoxLayout, QHBoxLayout, QWidget, QPushButton, QTextEdit, QListWidget, QLabel, QTabWidget
)
from PyQt5.QtGui import QPalette, QColor, QFont, QLinearGradient, QGradient
from PyQt5.QtCore import Qt
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
import docx
import PyPDF2
from docx import Document
import uuid
import re
from io import BytesIO

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("CV Analyzer Pro")
        self.setGeometry(100, 100, 1200, 800)

        # Load spaCy model
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load spaCy model: {str(e)}", QMessageBox.Ok)
            sys.exit(1)

        # Initialize data
        self.files = []
        self.results = []
        self.skills_data = {}
        self.candidate_colors = ['#0E6CFF', '#28A745', '#6F42C1', '#FF5733', '#FFC107', '#17A2B8', '#DC3545', '#6610F2']

        # Setup UI
        self.setup_ui()

    def setup_ui(self):
        # Set modern dark theme with vibrant accents
        palette = QPalette()
        palette.setColor(QPalette.Window, QColor("#1A1A2E"))
        palette.setColor(QPalette.WindowText, QColor("#F5F6FA"))
        palette.setColor(QPalette.Base, QColor("#2A2A3E"))
        palette.setColor(QPalette.Text, QColor("#F5F6FA"))
        palette.setColor(QPalette.Button, QColor("#0E6CFF"))
        palette.setColor(QPalette.ButtonText, QColor("#FFFFFF"))
        self.setPalette(palette)

        # Main widget and layout
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        main_layout = QHBoxLayout()
        main_layout.setSpacing(15)
        main_layout.setContentsMargins(15, 15, 15, 15)
        main_widget.setLayout(main_layout)

        # Left panel (controls)
        left_panel = QWidget()
        left_layout = QVBoxLayout()
        left_layout.setSpacing(10)
        left_panel.setLayout(left_layout)
        left_panel.setStyleSheet("background: qlineargradient(x1:0, y1:0, x2:1, y2:1, "
                               "stop:0 #1A1A2E, stop:1 #2A2A3E); border-radius: 10px; padding: 15px;")
        left_panel.setFixedWidth(400)

        # Title
        title_label = QLabel("CV Analyzer Pro")
        title_label.setFont(QFont("Segoe UI", 22, QFont.Bold))
        title_label.setStyleSheet("color: #00D4FF; margin-bottom: 20px;")
        left_layout.addWidget(title_label)

        # Job requirements input
        requirements_label = QLabel("Job Requirements")
        requirements_label.setFont(QFont("Segoe UI", 14, QFont.Bold))
        requirements_label.setStyleSheet("color: #F5F6FA;")
        left_layout.addWidget(requirements_label)

        self.requirements_input = QTextEdit()
        self.requirements_input.setFont(QFont("Segoe UI", 12))
        self.requirements_input.setStyleSheet("background: #2A2A3E; color: #F5F6FA; border: 1px solid #0E6CFF; border-radius: 5px; padding: 5px;")
        self.requirements_input.setFixedHeight(150)
        left_layout.addWidget(self.requirements_input)

        # Upload button
        self.upload_button = QPushButton("Upload CVs")
        self.upload_button.setFont(QFont("Segoe UI", 12, QFont.Bold))
        self.upload_button.setStyleSheet("""
            QPushButton {
                background-color: #0E6CFF;
                color: white;
                border-radius: 5px;
                padding: 12px;
            }
            QPushButton:hover {
                background-color: #00D4FF;
            }
        """)
        self.upload_button.clicked.connect(self.upload_files)
        left_layout.addWidget(self.upload_button)

        # File list
        self.file_list = QListWidget()
        self.file_list.setFont(QFont("Segoe UI", 12))
        self.file_list.setStyleSheet("background: #2A2A3E; color: #F5F6FA; border: 1px solid #0E6CFF; border-radius: 5px; padding: 5px;")
        left_layout.addWidget(self.file_list)

        # Analyze button
        self.analyze_button = QPushButton("Analyze CVs")
        self.analyze_button.setFont(QFont("Segoe UI", 12, QFont.Bold))
        self.analyze_button.setStyleSheet("""
            QPushButton {
                background-color: #6F42C1;
                color: white;
                border-radius: 5px;
                padding: 12px;
            }
            QPushButton:hover {
                background-color: #9F7AEA;
            }
        """)
        self.analyze_button.clicked.connect(self.analyze_cvs)
        left_layout.addWidget(self.analyze_button)

        # Generate report button
        self.generate_report_button = QPushButton("Generate Report")
        self.generate_report_button.setFont(QFont("Segoe UI", 12, QFont.Bold))
        self.generate_report_button.setStyleSheet("""
            QPushButton {
                background-color: #28A745;
                color: white;
                border-radius: 5px;
                padding: 12px;
            }
            QPushButton:hover {
                background-color: #38C759;
            }
        """)
        self.generate_report_button.clicked.connect(self.generate_report)
        left_layout.addWidget(self.generate_report_button)

        main_layout.addWidget(left_panel)

        # Right panel (results and graphs)
        right_panel = QWidget()
        right_layout = QVBoxLayout()
        right_layout.setSpacing(10)
        right_panel.setLayout(right_layout)
        right_panel.setStyleSheet("background: #16213E; border-radius: 10px; padding: 15px;")

        # Results table
        self.result_table = QTableWidget()
        self.result_table.setColumnCount(4)
        self.result_table.setHorizontalHeaderLabels(["Candidate", "Match Score", "Top Skills", "Best Fit"])
        self.result_table.setFont(QFont("Segoe UI", 12))
        self.result_table.setStyleSheet("""
            QTableWidget {
                background: #2A2A3E;
                color: #F5F6FA;
                gridline-color: #0E6CFF;
                selection-background-color: #0E6CFF;
            }
            QHeaderView::section {
                background-color: #0E6CFF;
                color: white;
                padding: 8px;
                font-weight: bold;
                font-size: 14px;
            }
        """)
        self.result_table.horizontalHeader().setStretchLastSection(True)
        right_layout.addWidget(self.result_table)

        # Tab widget for graphs
        self.tab_widget = QTabWidget()
        self.tab_widget.setFont(QFont("Segoe UI", 12))
        self.tab_widget.setStyleSheet("""
            QTabWidget::pane {
                border: 1px solid #0E6CFF;
                background: #16213E;
            }
            QTabBar::tab {
                background: #2A2A3E;
                color: #F5F6FA;
                padding: 10px;
                margin-right: 2px;
                border-top-left-radius: 5px;
                border-top-right-radius: 5px;
            }
            QTabBar::tab:selected {
                background: #0E6CFF;
                color: white;
            }
        """)
        right_layout.addWidget(self.tab_widget)

        # Summary graph tab
        self.summary_figure = plt.Figure(figsize=(6, 4), facecolor="#16213E")
        self.summary_canvas = FigureCanvas(self.summary_figure)
        self.summary_canvas.setStyleSheet("background: #16213E;")
        summary_tab = QWidget()
        summary_layout = QVBoxLayout()
        summary_layout.addWidget(self.summary_canvas)
        summary_tab.setLayout(summary_layout)
        self.tab_widget.addTab(summary_tab, "Summary")

        main_layout.addWidget(right_panel)

    def extract_text_from_file(self, filepath):
        try:
            if filepath.endswith('.pdf'):
                with open(filepath, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() or ""
            elif filepath.endswith('.docx'):
                doc = Document(filepath)
                text = "\n".join([para.text for para in doc.paragraphs])
            elif filepath.endswith('.txt'):
                with open(filepath, 'r', encoding='utf-8') as file:
                    text = file.read()
            else:
                return ""
            return text
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to read {filepath}: {str(e)}", QMessageBox.Ok)
            return ""

    def extract_skills(self, text):
        doc = self.nlp(text.lower())
        skills = []
        skill_keywords = [
            "python", "java", "javascript", "sql", "machine learning", "data analysis",
            "project management", "communication", "leadership", "cloud computing",
            "aws", "azure", "docker", "kubernetes", "react", "angular", "node.js",
            "typescript", "devops", "agile", "scrum", "database management"
        ]
        for token in doc:
            if token.text in skill_keywords:
                skills.append(token.text)
        for chunk in doc.noun_chunks:
            if chunk.text in skill_keywords:
                skills.append(chunk.text)
        return list(set(skills))

    def calculate_match_score(self, cv_text, job_requirements):
        job_doc = self.nlp(job_requirements.lower())
        cv_doc = self.nlp(cv_text.lower())
        similarity = job_doc.similarity(cv_doc) * 100
        keyword_matches = sum(1 for req in job_requirements.lower().split() if req in cv_text.lower())
        keyword_score = (keyword_matches / max(len(job_requirements.split()), 1)) * 50
        return min(round(similarity + keyword_score, 2), 100)

    def upload_files(self):
        try:
            files, _ = QFileDialog.getOpenFileNames(self, "Upload CVs", "", "CV Files (*.pdf *.docx *.txt)")
            if files:
                self.files = files
                self.file_list.clear()
                self.file_list.addItems([f.split('/')[-1] for f in files])
                QMessageBox.information(self, "Success", f"{len(files)} CVs uploaded!", QMessageBox.Ok)
                self.result_table.setRowCount(0)
                self.tab_widget.clear()
                self.tab_widget.addTab(QWidget(), "Summary")  # Placeholder until analysis
                self.summary_figure.clear()
                self.summary_canvas.draw()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to upload files: {str(e)}", QMessageBox.Ok)

    def analyze_cvs(self):
        try:
            if not self.files:
                QMessageBox.warning(self, "No Files", "Please upload CVs before analyzing.", QMessageBox.Ok)
                return

            job_requirements = self.requirements_input.toPlainText()
            if not job_requirements:
                QMessageBox.warning(self, "No Requirements", "Please provide job requirements.", QMessageBox.Ok)
                return

            self.results = []
            self.skills_data = {}

            for file in self.files:
                text = self.extract_text_from_file(file)
                if not text:
                    continue
                skills = self.extract_skills(text)
                score = self.calculate_match_score(text, job_requirements)
                self.skills_data[file.split('/')[-1]] = skills
                self.results.append((file.split('/')[-1], score, skills))

            if not self.results:
                QMessageBox.warning(self, "No Results", "No valid CVs were processed.", QMessageBox.Ok)
                return

            self.results.sort(key=lambda x: x[1], reverse=True)
            best_candidate = self.results[0][0] if self.results else ""
            self.display_results(best_candidate)
            self.display_graphs()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Analysis failed: {str(e)}", QMessageBox.Ok)

    def display_results(self, best_candidate):
        self.result_table.setRowCount(len(self.results))
        for i, (candidate, score, skills) in enumerate(self.results):
            self.result_table.setItem(i, 0, QTableWidgetItem(candidate))
            self.result_table.setItem(i, 1, QTableWidgetItem(f"{score}%"))
            self.result_table.setItem(i, 2, QTableWidgetItem(", ".join(skills[:5])))
            self.result_table.setItem(i, 3, QTableWidgetItem("Best" if candidate == best_candidate else ""))

    def display_graphs(self):
        # Clear existing tabs
        self.tab_widget.clear()

        # Summary tab
        self.summary_figure.clear()
        ax1 = self.summary_figure.add_subplot(121)
        candidates = [r[0] for r in self.results]
        scores = [r[1] for r in self.results]
        colors = self.candidate_colors[:len(candidates)]
        bars = ax1.bar(candidates, scores, color=colors)
        ax1.set_xlabel('Candidates', color='#F5F6FA', fontsize=12)
        ax1.set_ylabel('Match Score (%)', color='#F5F6FA', fontsize=12)
        ax1.set_title('Candidate Match Scores', color='#F5F6FA', fontsize=14)
        ax1.set_ylim(0, 100)
        ax1.grid(True, axis='y', linestyle='--', alpha=0.7)
        ax1.set_xticklabels(candidates, rotation=45, ha='right', color='#F5F6FA', fontsize=10)
        ax1.tick_params(colors='#F5F6FA')
        for bar in bars:
            yval = bar.get_height()
            ax1.text(bar.get_x() + bar.get_width()/2, yval + 2, f'{yval}%', ha='center', color='#F5F6FA', fontsize=10)

        ax2 = self.summary_figure.add_subplot(122)
        skill_counts = {}
        for _, _, skills in self.results:
            for skill in skills:
                skill_counts[skill] = skill_counts.get(skill, 0) + 1
        if skill_counts:
            skills, counts = zip(*skill_counts.items())
            ax2.pie(counts, labels=skills, autopct='%1.1f%%', textprops={'color': '#F5F6FA', 'fontsize': 10})
            ax2.set_title('Skill Distribution', color='#F5F6FA', fontsize=14)

        self.summary_figure.tight_layout()
        self.summary_canvas.draw()
        summary_tab = QWidget()
        summary_layout = QVBoxLayout()
        summary_layout.addWidget(self.summary_canvas)
        summary_tab.setLayout(summary_layout)
        self.tab_widget.addTab(summary_tab, "Summary")

        # Individual candidate tabs
        for i, (candidate, score, skills) in enumerate(self.results):
            figure = plt.Figure(figsize=(6, 4), facecolor="#16213E")
            canvas = FigureCanvas(figure)
            canvas.setStyleSheet("background: #16213E;")
            ax = figure.add_subplot(111)
            skill_counts = {skill: 1 for skill in skills[:5]}  # Show top 5 skills
            if skill_counts:
                skills, counts = zip(*skill_counts.items())
                ax.bar(skills, counts, color=self.candidate_colors[i % len(self.candidate_colors)])
                ax.set_title(f"{candidate} - Score: {score}%", color='#F5F6FA', fontsize=14)
                ax.set_ylabel('Skill Presence', color='#F5F6FA', fontsize=12)
                ax.set_xticklabels(skills, rotation=45, ha='right', color='#F5F6FA', fontsize=10)
                ax.tick_params(colors='#F5F6FA')
                ax.set_ylim(0, 2)
                ax.grid(True, axis='y', linestyle='--', alpha=0.7)
            figure.tight_layout()
            tab = QWidget()
            tab_layout = QVBoxLayout()
            tab_layout.addWidget(canvas)
            tab.setLayout(tab_layout)
            self.tab_widget.addTab(tab, candidate)

    def generate_report(self):
        try:
            if not self.results:
                QMessageBox.warning(self, "No Results", "Please analyze CVs before generating a report.", QMessageBox.Ok)
                return

            filepath, _ = QFileDialog.getSaveFileName(self, "Save Report", "", "PDF Files (*.pdf);;Word Files (*.docx)")
            if filepath:
                if filepath.endswith('.pdf'):
                    self.generate_pdf_report(filepath)
                else:
                    self.generate_docx_report(filepath)
                QMessageBox.information(self, "Success", "Report generated successfully!", QMessageBox.Ok)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate report: {str(e)}", QMessageBox.Ok)

    def generate_pdf_report(self, filepath):
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []

        elements.append(Paragraph("CV Analysis Report", styles['Title']))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph("Analysis Summary", styles['Heading2']))
        elements.append(Paragraph(f"Total CVs analyzed: {len(self.results)}", styles['Normal']))
        elements.append(Paragraph(f"Best Candidate: {self.results[0][0]} (Score: {self.results[0][1]}%)", styles['Normal']))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph("Detailed Results", styles['Heading2']))

        data = [["Candidate", "Match Score", "Top Skills"]]
        for candidate, score, skills in self.results:
            data.append([candidate, f"{score}%", ", ".join(skills[:5])])

        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#0E6CFF")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor("#E9ECEF")),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        elements.append(table)

        doc.build(elements)

    def generate_docx_report(self, filepath):
        doc = Document()
        doc.add_heading('CV Analysis Report', 0)
        doc.add_heading('Analysis Summary', level=2)
        doc.add_paragraph(f'Total CVs analyzed: {len(self.results)}')
        doc.add_paragraph(f'Best Candidate: {self.results[0][0]} (Score: {self.results[0][1]}%)')
        doc.add_heading('Detailed Results', level=2)

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Candidate'
        hdr_cells[1].text = 'Match Score'
        hdr_cells[2].text = 'Top Skills'

        for candidate, score, skills in self.results:
            row_cells = table.add_row().cells
            row_cells[0].text = candidate
            row_cells[1].text = f'{score}%'
            row_cells[2].text = ', '.join(skills[:5])

        doc.save(filepath)

if __name__ == "__main__":
    try:
        app = QApplication(sys.argv)
        window = MainWindow()
        window.show()
        sys.exit(app.exec_())
    except Exception as e:
        print(f"Application failed: {str(e)}")
        sys.exit(1)